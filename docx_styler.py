# docx_styler.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn, nsdecls

class DocxStyler:
    """Class to handle styling of Word documents"""
    
    def __init__(self, doc=None):
        """Initialize with an existing document or create a new one"""
        self.doc = doc if doc else Document()
    
    def apply_heading_style(self, heading, level=1):
        """Apply a consistent style to a heading"""
        heading_paragraph = self.doc.add_heading(heading, level=level)
        
        # Style settings based on heading level
        if level == 1:
            heading_paragraph.style.font.size = Pt(18)
            heading_paragraph.style.font.bold = True
            heading_paragraph.style.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
        elif level == 2:
            heading_paragraph.style.font.size = Pt(14)
            heading_paragraph.style.font.bold = True
            heading_paragraph.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        else:
            heading_paragraph.style.font.size = Pt(12)
            heading_paragraph.style.font.bold = True
        
        return heading_paragraph
    
    def add_styled_paragraph(self, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, size=None):
        """Add a paragraph with custom styling"""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = alignment
        
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        
        if size:
            run.font.size = Pt(size)
        
        return paragraph
    
    def add_styled_table(self, data, style='Table Grid', has_header=True):
        """Add a table with custom styling"""
        if not data:
            return None
        
        # Determine table dimensions
        num_rows = len(data)
        num_cols = len(data[0]) if num_rows > 0 else 0
        
        if num_rows == 0 or num_cols == 0:
            return None
        
        # Create the table
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Populate the table with data
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # Style header row if present
                if has_header and i == 0:
                    # Make header bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    
                    # Set header row background color (light gray)
                    shading_xml = f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>'
                    cell._element.get_or_add_tcPr().append(parse_xml(shading_xml))
        
        return table
    
    def add_table_from_dataframe(self, df, style='Table Grid', has_header=True):
        """Add a table from a pandas DataFrame"""
        # Convert DataFrame to list of lists
        if has_header:
            data = [df.columns.tolist()] + df.values.tolist()
        else:
            data = df.values.tolist()
        
        return self.add_styled_table(data, style, has_header)
    
    def add_section_break(self):
        """Add a section break to the document"""
        self.doc.add_section()
    
    def add_page_break(self):
        """Add a page break to the document"""
        self.doc.add_page_break()
    
    def get_document(self):
        """Get the styled document"""
        return self.doc
## Helper function to parse XML
